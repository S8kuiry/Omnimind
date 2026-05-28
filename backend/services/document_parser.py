import fitz  # PyMuPDF  — for PDF
import io

# ── PDF ────────────────────────────────────────────────────────────

def extract_text_from_pdf(file_bytes: bytes) -> list[dict]:
    """
    Extracts text page-by-page from a PDF.
    Returns: [{ "page": int, "text": str }, ...]
    """
    pages = []
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text()
        if text.strip():
            pages.append({
                "page": page_num + 1,
                "text": text.strip()
            })
    doc.close()
    return pages


# ── DOCX ───────────────────────────────────────────────────────────

def extract_text_from_docx(file_bytes: bytes, paragraphs_per_page: int = 15) -> list[dict]:
    """
    Extracts text from a DOCX file.
    Word docs have no hard page boundaries in the text layer,
    so we simulate pages by grouping N paragraphs together.
    This keeps the page/citation UX consistent with PDF.

    Returns: [{ "page": int, "text": str }, ...]
    """
    try:
        from docx import Document  # python-docx
    except ModuleNotFoundError as exc:
        raise ModuleNotFoundError(
            "Missing dependency for DOCX support. Install `python-docx` in the same "
            "environment that runs the FastAPI server."
        ) from exc

    doc = Document(io.BytesIO(file_bytes))

    # collect non-empty paragraphs
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # also extract text from tables (often important in docs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    paragraphs.append(cell_text)

    if not paragraphs:
        return []

    # group into simulated pages
    pages = []
    for i in range(0, len(paragraphs), paragraphs_per_page):
        group = paragraphs[i:i + paragraphs_per_page]
        page_text = "\n\n".join(group)
        if page_text.strip():
            pages.append({
                "page": (i // paragraphs_per_page) + 1,
                "text": page_text.strip()
            })

    return pages


# ── Unified entry point ────────────────────────────────────────────

def extract_text_from_document(file_bytes: bytes, filename: str) -> list[dict]:
    """
    Routes to the correct extractor based on file extension.
    Use this in your upload endpoint instead of calling
    extract_text_from_pdf directly.

    Raises ValueError for unsupported file types.
    """
    name = filename.lower()

    if name.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)

    if name.endswith(".docx"):
        return extract_text_from_docx(file_bytes)

    raise ValueError(f"Unsupported file type: {filename}. Only PDF and DOCX are supported.")